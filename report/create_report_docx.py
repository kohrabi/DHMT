from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, Inches


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "report" / "bao_cao_do_an.docx"
SCREENSHOTS = ROOT / "screenshoots"


def set_run_font(run, size=13, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_paragraph(doc, text="", align=None, bold=False, italic=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 14, bold=True)
    if not p.runs:
        r = p.add_run(text)
        set_run_font(r, size=16 if level == 1 else 14, bold=True)
    else:
        p.runs[0].text = text
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[i].paragraphs[0]
            if i == 0 or (headers[i].lower() == "mssv"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_run_font(r)
    return table


def add_figure(doc, filename, caption):
    image_path = SCREENSHOTS / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(image_path), width=Inches(6.3))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, size=12, italic=True)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"].font.size = Pt(13)

add_paragraph(doc, "ĐẠI HỌC QUỐC GIA TP. HỒ CHÍ MINH", WD_ALIGN_PARAGRAPH.CENTER, bold=True)
add_paragraph(doc, "TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN", WD_ALIGN_PARAGRAPH.CENTER, bold=True)
doc.add_paragraph()
add_paragraph(doc, "MÔN HỌC: ĐỒ HỌA MÁY TÍNH - CS105", WD_ALIGN_PARAGRAPH.CENTER, bold=True)
doc.add_paragraph()
title = add_paragraph(doc, "BÁO CÁO ĐỒ ÁN", WD_ALIGN_PARAGRAPH.CENTER, bold=True)
title.runs[0].font.size = Pt(20)
title2 = add_paragraph(doc, "MÔ PHỎNG GAME SUPER MARIO 3D", WD_ALIGN_PARAGRAPH.CENTER, bold=True)
title2.runs[0].font.size = Pt(20)
add_paragraph(doc, "Sử dụng Three.js và Rapier Physics", WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_table(
    doc,
    ["Thông tin", "Nội dung"],
    [
        ["Giảng viên hướng dẫn", "Cáp Phạm Đình Thăng"],
        ["Lớp", "CS105.Q22"],
        ["Học kỳ", "2025-2026"],
    ],
)
doc.add_paragraph()
add_paragraph(doc, "TP. Hồ Chí Minh, 2026", WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

add_heading(doc, "1. Giới thiệu nhóm")
add_table(
    doc,
    ["STT", "Họ và tên", "MSSV", "Vai trò / đóng góp"],
    [
        ["1", "Nguyễn Mỹ Thống", "23521527", "Xây dựng scene, hệ thống đồ họa, ánh sáng, hiệu ứng, tài nguyên 3D và trình bày báo cáo."],
        ["2", "Đặng Trần Anh Hào", "23520444", "Xây dựng gameplay, tích hợp engine, xử lý nhân vật, vật lý va chạm và hoàn thiện demo."],
    ],
)

add_heading(doc, "2. Giới thiệu đề tài")
add_paragraph(doc, "Đồ án xây dựng một chương trình mô phỏng game platformer 3D lấy cảm hứng từ Super Mario. Người chơi điều khiển nhân vật di chuyển, nhảy qua chướng ngại vật, tương tác với các khối câu hỏi, thu thập vật phẩm và tiêu diệt kẻ địch trong môi trường 3D phong cách low-poly.")
add_paragraph(doc, "Dự án được phát triển trên nền WebGL thông qua thư viện Three.js, kết hợp Rapier3D để xử lý vật lý và va chạm. Mục tiêu chính là vận dụng các kiến thức đồ họa máy tính như phép chiếu phối cảnh, phép biến đổi affine, ánh sáng, bóng đổ, texture, mô hình 3D và animation vào một sản phẩm có tính tương tác thực tế.")

add_heading(doc, "3. Công nghệ sử dụng")
add_table(
    doc,
    ["Công nghệ", "Vai trò trong đồ án"],
    [
        ["Three.js", "Dựng cảnh 3D, camera, ánh sáng, vật liệu, model GLTF, shadow mapping và render WebGL."],
        ["Rapier3D", "Mô phỏng vật lý, rigid body, collider, va chạm nhân vật, enemy, block, platform và kill zone."],
        ["Vite / Bun", "Môi trường phát triển và chạy demo cục bộ."],
        ["Tween.js", "Hỗ trợ nội suy chuyển động, hiệu ứng UI, scale và animation phụ."],
        ["GLTF Assets", "Nạp mô hình 3D cho nhân vật, map, enemy, coin, block và các vật thể trang trí."],
    ],
)

add_heading(doc, "4. Mô tả các chức năng trong đồ án")
add_heading(doc, "4.1. Các chức năng đã làm được", level=2)
add_bullets(doc, [
    "Main Menu: Có màn hình menu chính với nền parallax, chữ 3D SUPER MARIO, nút Start Game và Settings.",
    "Điều khiển nhân vật: Người chơi có thể di chuyển trái/phải, nhảy, chạy, ngồi và tương tác với môi trường.",
    "Camera 2.5D: Camera phối cảnh bám theo nhân vật, giới hạn trong vùng chơi và hỗ trợ OrbitControls để debug.",
    "Môi trường màn chơi: Có nền đồi xanh, cây, hoa, hàng rào, đá, ống nước, nền trời gradient và các platform.",
    "Question Block và Brick: Người chơi có thể tương tác với block để tạo hiệu ứng bật, score popup hoặc vật phẩm.",
    "Coin: Đồng xu có hiệu ứng xoay và dùng phép biến đổi trong không gian 3D.",
    "Enemy: Có Goomba/Koopa với logic di chuyển, va chạm và trạng thái bị tiêu diệt.",
    "Power-up: Có vật phẩm nấm và hiệu ứng phóng to nhân vật khi ăn power-up.",
    "Score Popup: Hiển thị điểm số dạng chữ 3D khi người chơi tiêu diệt enemy hoặc tương tác với object.",
    "Hiệu ứng hạt: Có hiệu ứng bụi khi nhân vật chạy/nhảy/chạm đất, giúp chuyển động sinh động hơn.",
    "Ánh sáng và bóng đổ: Sử dụng DirectionalLight, HemisphereLight và PCFSoft Shadow Map để tạo chiều sâu cho cảnh.",
    "Physics và collision: Tích hợp Rapier3D cho hệ thống collider, rigid body, one-way platform, kill zone và kiểm tra va chạm.",
    "Kiến trúc engine: Tách riêng engine và game logic gồm SceneManager, GameObject, ContentManager, InputManager, PhysicsWorld, Animator và GameTimer.",
])

add_heading(doc, "4.2. Các chức năng mới / điểm nổi bật", level=2)
add_bullets(doc, [
    "3D Outline Text: Tạo chữ điểm số nổi bật bằng kỹ thuật double-mesh, giúp chữ có viền tương phản khi hiển thị trong cảnh 3D.",
    "Viewport Clamping: Camera được giới hạn động để không vượt ra ngoài PlayZone.",
    "Dynamic Shadow Tracking: Nguồn sáng và vùng đổ bóng bám theo camera để duy trì chất lượng bóng khi người chơi di chuyển.",
    "Z-up to Y-up Pipeline: Chuyển đổi hệ tọa độ model từ Blender sang Three.js để đồng bộ asset 3D.",
    "One-way Platform: Platform cho phép nhân vật nhảy xuyên từ dưới lên và đứng lại khi tiếp đất từ phía trên.",
    "Subpixel Movement Scale: Tối ưu cảm giác điều khiển và chuyển động của nhân vật theo phong cách platformer.",
])

add_heading(doc, "4.3. Các chức năng chưa làm được / hạn chế", level=2)
add_bullets(doc, [
    "Chưa có nhiều màn chơi, hiện mới tập trung vào một level demo.",
    "Việc thiết kế bản đồ hiện được thực hiện bằng Blender và import vào dự án, chưa có công cụ chỉnh level trực tiếp ngay trong game.",
    "Chưa hoàn thiện toàn bộ HUD như điểm, mạng, thời gian và trạng thái nhân vật.",
    "Chưa tích hợp âm thanh nền và hiệu ứng âm thanh.",
    "Một số power-up nâng cao như Racoon chưa được hoàn thiện đầy đủ.",
    "Chưa tối ưu hoàn chỉnh cho thiết bị di động.",
    "Chưa có hệ thống lưu điểm cao hoặc save/load tiến trình.",
])

add_heading(doc, "5. Show các chức năng đã làm")
add_paragraph(doc, "Các hình bên dưới minh họa những chức năng chính đã hoàn thiện trong đồ án.")
figures = [
    ("menu.png", "Hình 1. Màn hình Main Menu với chữ 3D, nền parallax, nút Start Game và Settings."),
    ("gameplay.png", "Hình 2. Màn chơi chính với nhân vật, Question Block, platform, pipe, nền đồi xanh và bóng đổ."),
    ("shadow_lightning.png", "Hình 3. Hệ thống chiếu sáng và bóng đổ trong môi trường 3D."),
    ("score_popup.png", "Hình 4. Hiệu ứng Score Popup khi nhân vật tương tác với block hoặc enemy."),
    ("death.png", "Hình 5. Hiệu ứng nhân vật/enemy bị tiêu diệt và trạng thái death spin trong gameplay."),
]
for filename, caption in figures:
    add_figure(doc, filename, caption)

add_heading(doc, "6. Kết luận")
add_paragraph(doc, "Đồ án đã hoàn thành một sản phẩm game platformer 3D có đầy đủ các thành phần cơ bản: menu, nhân vật điều khiển được, camera, môi trường 3D, enemy, vật phẩm, block tương tác, score popup, hiệu ứng hạt, ánh sáng, bóng đổ và hệ thống vật lý. Thông qua đồ án, nhóm đã áp dụng được nhiều kiến thức quan trọng của môn Đồ họa máy tính vào một chương trình có khả năng chạy và tương tác thực tế.")
add_paragraph(doc, "Trong tương lai, dự án có thể tiếp tục phát triển thêm nhiều level, âm thanh, HUD, công cụ chỉnh level trực tiếp trong game, lưu điểm cao, tối ưu mobile và các power-up nâng cao để hoàn thiện trải nghiệm giống một game platformer hoàn chỉnh hơn.")

doc.save(OUT)
print(OUT)
